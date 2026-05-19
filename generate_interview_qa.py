"""
Generate ML Project Interview Q&A Word Document for Pamela Austin
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

def add_heading(doc, text, level=1, color=None):
    heading = doc.add_heading(text, level=level)
    heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = heading.runs[0]
    if color:
        run.font.color.rgb = RGBColor(*color)
    run.font.bold = True
    return heading


def add_label(doc, label, text, label_color=(30, 100, 170)):
    """Add a Q: or A: paragraph with label in bold color."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(4)
    lbl = p.add_run(label + " ")
    lbl.bold = True
    lbl.font.color.rgb = RGBColor(*label_color)
    lbl.font.size = Pt(11)
    body = p.add_run(text)
    body.font.size = Pt(11)
    return p


def add_bullet_list(doc, items, indent=0):
    for item in items:
        p = doc.add_paragraph(item, style='List Bullet')
        p.paragraph_format.left_indent = Inches(0.25 + indent * 0.25)
        p.paragraph_format.space_after = Pt(2)
        for run in p.runs:
            run.font.size = Pt(11)


def add_section_break(doc):
    doc.add_paragraph()


def set_doc_margins(doc):
    section = doc.sections[0]
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)


def build_document():
    doc = Document()
    set_doc_margins(doc)

    # ── Title Page ──────────────────────────────────────────────────────────
    title = doc.add_heading("ML Project Interview Q&A", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].font.color.rgb = RGBColor(30, 100, 170)

    subtitle = doc.add_paragraph("Pamela Austin | Data Science & Analytics Portfolio")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.runs[0].font.size = Pt(13)
    subtitle.runs[0].font.color.rgb = RGBColor(90, 90, 90)
    doc.add_paragraph()

    note = doc.add_paragraph(
        "This document is organized as a conversational interview guide. Each project "
        "section leads with context, then flows through follow-up questions as an "
        "interviewer would naturally progress. Use this to practice storytelling as "
        "well as technical depth."
    )
    note.runs[0].font.size = Pt(10)
    note.runs[0].font.color.rgb = RGBColor(100, 100, 100)
    note.runs[0].italic = True
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # PROJECT 1 – AT&T Predictive ML Model
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "1. PREDICTIVE ML FORECASTING – AT&T", level=1, color=(30, 100, 170))
    p = doc.add_paragraph()
    p.add_run("89% Timeline Forecast Accuracy  |  28% Reduction in Project Overruns  |  "
              "30% Reduction in Escalations  |  March 2024 – March 2025").bold = True

    add_label(doc, "Q:",
              "Tell me about the predictive forecasting model you built at AT&T. "
              "What was the business problem you were solving?")

    add_label(doc, "A:",
              "AT&T's Marketing Operations division was managing 200+ concurrent campaigns and had "
              "no systematic way to predict which projects were at risk of overrunning their timelines "
              "before it happened. Project managers were relying on subjective gut-feel estimates. "
              "Similar projects repeatedly escalated to leadership for the same root causes, but there "
              "was no learning loop from historical outcomes. Leadership only found out a project was "
              "in trouble after it was already delayed—at which point your options for recovery are "
              "very limited.\n\n"
              "I sat in a budget meeting early on where the CFO showed a slide: 'Q3 forecast missed "
              "by $47M.' The room went quiet. That miss meant delayed infrastructure investments and "
              "hiring freezes in two regions. That moment made it concrete—forecasting isn't an "
              "academic exercise. Every percentage point of accuracy translates into real resource "
              "allocation decisions.\n\n"
              "My objective was to build a forecasting system that predicted project completion "
              "timelines and revenue, flagged at-risk campaigns proactively, and surfaced those "
              "insights to project managers through their existing Power BI dashboards.")

    add_label(doc, "Q:",
              "Walk me through your data sources and how you brought them together.")

    add_label(doc, "A:",
              "Data was scattered across four source systems with no unified view: Salesforce CRM "
              "for contract and opportunity data, SAP ERP for billing and revenue transactions, "
              "Oracle Financials for general ledger and actuals, and Workfront for project timelines, "
              "resource allocation, and task-level data.\n\n"
              "We centralized everything into a Snowflake data warehouse. Fivetran handled the SaaS "
              "connectors for real-time sync, AWS Glue ETL managed the heavier extraction jobs, and "
              "dbt built the transformation layer—staging models to clean and standardize, intermediate "
              "models for business logic, and mart-layer fact and dimension tables that the forecasting "
              "pipeline consumed. The mart layer gave us conformed dimensions for employee, campaign, "
              "and time hierarchies that all three model components shared.\n\n"
              "By the time data reached the ML feature engineering step, it was already clean, "
              "historized with SCD Type 2 for slowly changing dimensions, and validated through "
              "dbt's automated test suite.")

    add_label(doc, "Q:", "Tell me about the feature engineering. What went into the model?")

    add_label(doc, "A:",
              "This was honestly where most of the performance gain came from. I initially plateaued "
              "at around 82% accuracy regardless of which algorithm I tried. The breakthrough came "
              "from going back to domain fundamentals—spending time with data engineers and finance "
              "stakeholders to understand what actually drives revenue and timeline patterns at AT&T.\n\n"
              "I ended up engineering 87 features. The main categories were:\n")
    add_bullet_list(doc, [
        "Lag features: revenue and project metrics at 7, 30, 90, and 365-day lags—matching weekly, monthly, quarterly, and annual business cycles",
        "Rolling statistics: rolling means and standard deviations over 7, 30, and 90-day windows to capture trend and volatility",
        "Seasonality features: day of week, month, quarter, month-end and quarter-end flags (enterprise spending spikes at period close)",
        "Growth and momentum features: month-over-month and year-over-year revenue growth rates, plus revenue acceleration (change in growth rate)",
        "Interaction terms: utilization × capacity, GDP growth × consumer confidence index, revenue per active customer",
        "Holiday proximity: days to nearest major holiday (Black Friday through New Year's drives enterprise spending patterns)",
    ])
    p = doc.add_paragraph(
        "That feature work—not the algorithm—jumped accuracy from 82% to 91% before I even tuned "
        "the model ensemble. It reinforced a lesson I now apply everywhere: domain knowledge beats "
        "platform complexity every time."
    )
    p.runs[0].font.size = Pt(11)

    add_label(doc, "Q:",
              "You mentioned a model ensemble. Walk me through the architecture—"
              "what components made up the final system?")

    add_label(doc, "A:",
              "The final deployed model was a weighted ensemble of three components, each contributing "
              "different strengths:\n\n"
              "SQL Trend Analysis (35% weight): Snowflake stored procedures computing exponentially "
              "weighted moving averages, linear trend projections using REGR_SLOPE, and seasonal "
              "adjustment factors based on same-period prior year comparisons. The strength here is "
              "speed, transparency, and tight integration with the data warehouse—any analyst can "
              "read the SQL and understand exactly what it's doing.\n\n"
              "Python Gradient Boosting Machine (40% weight): scikit-learn's "
              "GradientBoostingRegressor trained on the 87-feature matrix. GBM captures non-linear "
              "relationships that SQL can't—for example, the interaction between an economic "
              "downturn and a network utilization drop affects revenue in ways that a moving average "
              "misses entirely. Hyperparameters: 200 estimators, max depth 6, learning rate 0.05, "
              "subsample 0.8, feature randomization via max_features='sqrt'.\n\n"
              "Power BI DAX Calculations (25% weight): Business logic adjustments for known "
              "one-time events, executive overrides for strategic initiatives, and AT&T policy "
              "floors (forecasts can't be below 95% of prior year). This layer gave the Finance "
              "team a way to apply their judgment without breaking the statistical foundation.\n\n"
              "Tested approaches and their error rates:")
    add_bullet_list(doc, [
        "Legacy Excel moving average: 29.0% MAPE (baseline)",
        "SQL moving averages: 18.0% MAPE",
        "Power BI DAX only: 14.0% MAPE",
        "Python GBM only: 12.0% MAPE",
        "Weighted ensemble (deployed): 5.8% MAPE  ←  R² = 0.967",
    ])

    add_label(doc, "Q:",
              "How did you handle cross-validation for a time series problem? "
              "What was your evaluation approach?")

    add_label(doc, "A:",
              "Time series data requires a fundamentally different validation approach than "
              "typical tabular ML—you cannot randomly shuffle your data and do standard k-fold "
              "cross-validation because that creates data leakage. If a test fold contains data "
              "from before the training fold, the model will appear more accurate than it actually "
              "is in production.\n\n"
              "I used TimeSeriesSplit from scikit-learn—5 splits with a 30-day gap between each "
              "training period and its test window. The gap is important: it prevents the model "
              "from using recent data that's adjacent to the forecast period, which simulates "
              "real-world conditions where you're always forecasting forward in time.\n\n"
              "Primary evaluation metrics were MAPE (Mean Absolute Percentage Error) for "
              "interpretability—5.8% MAPE means the model's forecast is within 5.8% of actual "
              "revenue on average, which executives immediately understand—and R² (0.967) to "
              "measure overall explanatory power. I deliberately avoided using accuracy as a "
              "primary metric because this is a regression problem predicting continuous values, "
              "not a classification problem. The 89% 'accuracy' headline represents how often the "
              "forecast direction (up/down) and magnitude were within acceptable tolerance bands "
              "defined by the Finance team.")

    add_label(doc, "Q:",
              "You mentioned SHAP values. How did you use them, and what was the "
              "stakeholder impact?")

    add_label(doc, "A:",
              "SHAP (SHapley Additive exPlanations) was actually the difference between the project "
              "succeeding and being abandoned.\n\n"
              "After I built the ensemble and showed 89% accuracy results to the Finance VP, I got "
              "silence. Then: 'How do I know this isn't just luck? How do I explain to the board "
              "why we should trust a machine over our analysts?' That question nearly derailed "
              "everything. Finance teams had built their careers on manual forecasting. They weren't "
              "going to hand over decision-making to a black box.\n\n"
              "I spent two weeks implementing SHAP using TreeExplainer on the GBM component. For "
              "every forecast, SHAP decomposed the prediction into feature-level contributions: "
              "'Q4 revenue is up because holiday season (+$12M), contract renewals (+$8M), "
              "seasonal network maintenance (-$3M).' I built a 'Forecast Story' feature in "
              "Tableau that translated statistical outputs into plain English narratives alongside "
              "the SHAP waterfall charts.\n\n"
              "The turning point: a Finance Director said, 'This is exactly what I do in my head—"
              "but you've automated it and made it consistent.' Executive dashboard adoption went "
              "from 40% to 100% after adding the SHAP explainability layer. The technical accuracy "
              "was table stakes—the explainability was what actually drove adoption.")

    add_label(doc, "Q:",
              "How was the model deployed and how did you handle retraining?")

    add_label(doc, "A:",
              "The model was deployed through Azure ML Studio with automated retraining managed "
              "by Apache Airflow DAGs. MLflow tracked every experiment—hyperparameters, feature "
              "sets, cross-validation scores—so I had a full audit trail of model versions and "
              "could roll back if a retrained model underperformed.\n\n"
              "The Airflow DAG ran weekly: pull new actuals from Snowflake, append to the "
              "training dataset, retrain the GBM component, validate that new MAPE was within "
              "10% of the prior model's performance, and if validation passed, promote to "
              "production automatically. If validation failed, it triggered an alert for "
              "manual review before any production change.\n\n"
              "The SQL and DAX components updated automatically because they ran directly against "
              "the live Snowflake data—no retraining needed. That's one of the advantages of "
              "the ensemble architecture: the statistical components are always current, which "
              "stabilizes the ensemble even in periods when the GBM component hasn't retrained yet.\n\n"
              "Over the deployment period, the ensemble reduced project overruns by 28% and "
              "escalations to leadership by 30%, against the prior 12-month baseline.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # PROJECT 2 – AT&T Cloud Migration ETL
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "2. ENTERPRISE ETL PIPELINE – AT&T CLOUD MIGRATION", level=1, color=(30, 100, 170))
    p = doc.add_paragraph()
    p.add_run("78% Query Performance Improvement  |  87.5% Faster Processing  |  $250K Annual Savings").bold = True

    add_label(doc, "Q:",
              "You led a cloud migration that consolidated three enterprise systems into Snowflake. "
              "Walk me through the business problem and why this was necessary.")

    add_label(doc, "A:",
              "AT&T had three critical enterprise platforms operating as independent data silos—WebPhone "
              "(our customer interaction and call management platform logging 850K records daily), Bynder "
              "(digital asset management with 1.2M asset metadata entries), and Workfront (project and "
              "resource management). None of them talked to each other.\n\n"
              "Marketing analysts needed data from all three to answer basic questions: 'Which campaigns "
              "drove inbound calls? What assets were used in which projects? Where are we over-utilizing "
              "team capacity?' Answering those required pulling separate exports from three systems, "
              "manually joining them in Excel, and hoping the IDs matched. That process took 2–3 days "
              "and introduced data quality errors. Leadership wanted real-time analytics for 5,000+ users "
              "and we were nowhere near that. The migration to Snowflake was the solution.")

    add_label(doc, "Q:",
              "Tell me about the ETL architecture you designed. How did you approach extracting "
              "from three different source systems?")

    add_label(doc, "A:",
              "Each source system had different extraction patterns, so I designed three distinct "
              "ingestion tracks that fed into a unified Snowflake raw layer.\n\n"
              "For WebPhone, we used a REST API connector with incremental extraction based on "
              "timestamp watermarking—pulling new and updated call records every 15 minutes into "
              "Snowflake staging tables. This handled the high-velocity, high-volume nature of call logs.\n\n"
              "For Bynder DAM, we used Fivetran's pre-built connector since the asset metadata was "
              "lower volume but schema-sensitive. Fivetran handled schema drift automatically, which "
              "saved us from brittle custom code when Bynder released API updates.\n\n"
              "For Workfront, we used a combination of Workfront's REST API and their bulk export "
              "capability. Project data is naturally hierarchical—portfolios, programs, projects, tasks—"
              "so we modeled it into a normalized dimensional schema in Snowflake.\n\n"
              "All three feeds landed in immutable raw staging schemas. From there, dbt handled all "
              "transformations through staging → intermediate → mart layers with full lineage tracking.")

    add_label(doc, "Q:",
              "How did you achieve the 78% query performance improvement? What specific "
              "optimizations did you implement?")

    add_label(doc, "A:",
              "The performance gains came from a combination of Snowflake-native features and smart "
              "data modeling choices.\n\n"
              "First, we redesigned the data model. The source systems used transaction-style schemas "
              "with heavy normalization. We converted to a star schema in the mart layer—fact tables "
              "for campaign events, call records, and project milestones, with shared dimension tables "
              "for customers, campaigns, agents, and assets. This eliminated multi-table joins on the "
              "BI side.\n\n"
              "Second, we implemented Snowflake automatic clustering on the highest-cardinality fact "
              "tables using date and campaign_id as clustering keys. This reduced partition pruning scan "
              "from full-table to targeted micro-partition reads.\n\n"
              "Third, we materialized the most commonly queried aggregations as Snowflake materialized "
              "views—campaign summary stats, weekly agent performance rollups, asset utilization rates. "
              "Queries that previously scanned 180M rows now hit pre-aggregated views.\n\n"
              "Baseline query time dropped from 6 hours for the full overnight load to 45 minutes, and "
              "ad-hoc query latency went from 45 seconds to under 3 seconds for the executive dashboards.")

    add_label(doc, "Q:",
              "What was the biggest technical challenge during the migration, and how did you "
              "solve it?")

    add_label(doc, "A:",
              "Data quality reconciliation across three systems with inconsistent identifier schemas. "
              "WebPhone used its own customer_id format, Workfront used account numbers from Salesforce, "
              "and Bynder had its own asset tagging taxonomy that only loosely mapped to campaign names "
              "in Workfront.\n\n"
              "The specific problem: we needed to join call records from WebPhone to campaigns in "
              "Workfront to understand which campaigns were driving inbound call volume. But there was "
              "no shared key—WebPhone had advertiser phone numbers, Workfront had Salesforce account IDs, "
              "and the link between them lived in a manually maintained Excel lookup table that was "
              "25% out of date.\n\n"
              "I built a probabilistic matching pipeline in Python that used phone number normalization, "
              "fuzzy account name matching (RapidFuzz library), and geolocation-based disambiguation to "
              "generate confidence-scored account linkages. Matches above 95% confidence were auto-accepted, "
              "those between 70–95% went to a human review queue, and anything below 70% was flagged for "
              "source system remediation. This reduced unmatched records from 23% to under 2%.\n\n"
              "The 2% residual was actually valuable—it surfaced orphaned WebPhone accounts that had "
              "never been properly provisioned in Workfront, which the operations team then cleaned up.")

    add_label(doc, "Q:",
              "How did you ensure zero-downtime migration for 5,000 concurrent users?")

    add_label(doc, "A:",
              "We used a phased cutover approach with parallel running periods. The strategy had four stages:\n\n"
              "First, we ran dual-write for 4 weeks—new data was written to both legacy systems and Snowflake "
              "simultaneously. This validated pipeline reliability without any user exposure.\n\n"
              "Second, we migrated read traffic incrementally by team. We started with the marketing analytics "
              "team (20 users) who were already comfortable with Snowflake queries. Their positive feedback "
              "and the absence of data discrepancy escalations gave us confidence to expand.\n\n"
              "Third, we built a reconciliation dashboard in Power BI that compared legacy system output "
              "vs. Snowflake output daily. Any metric that differed by more than 0.1% triggered an alert "
              "and a root-cause analysis.\n\n"
              "Fourth, the final cutover weekend involved a 4-hour maintenance window where we ran final "
              "delta loads, validated row counts across all 47 tables, and flipped DNS for the reporting "
              "layer. We had a 30-minute rollback plan ready and never needed it. System uptime went from "
              "98.5% to 99.95% post-migration.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # PROJECT 3 – AT&T Marketing Analytics Platform
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "3. MARKETING ANALYTICS PLATFORM – AT&T", level=1, color=(30, 100, 170))
    p = doc.add_paragraph()
    p.add_run("40% Reduction in Manual Work  |  150+ Self-Service Users  |  60% Fewer IT Data Requests").bold = True

    add_label(doc, "Q:",
              "You mentioned building an analytics platform that enabled 150+ self-service users. "
              "What was the environment like before you built it, and what problem were you solving?")

    add_label(doc, "A:",
              "When I joined, there were 3 analysts servicing the data needs of the entire marketing "
              "operations organization. Marketing managers were waiting 5–7 days for campaign performance "
              "reports. Worse, 27% of those reports contained errors—copy-paste mistakes, broken Excel "
              "formulas, stale data. The team was spending 15–20 hours a week just exporting data from "
              "Workfront, Salesforce, and ZoomInfo into spreadsheets.\n\n"
              "I remember my first stakeholder meeting—a campaign manager pulled out a USB drive with "
              "47 Excel files and called it her 'reporting system.' That moment made the problem concrete: "
              "brilliant marketing professionals were trapped doing data entry instead of marketing analysis. "
              "We were processing 500K+ marketing events monthly across 200+ campaigns and managing it "
              "through spreadsheets.")

    add_label(doc, "Q:",
              "How did you design the data model and what was the semantic layer architecture?")

    add_label(doc, "A:",
              "Before writing a single line of code, I spent a week shadowing each analyst through their "
              "Monday morning data ritual—watching someone spend 3 hours copy-pasting phone call data "
              "into a pivot table told me more about requirements than any formal spec.\n\n"
              "The architecture was a three-layer model in Snowflake:\n\n"
              "Raw layer: Fivetran-managed connectors pulling from Workfront, Salesforce, and ZoomInfo "
              "into immutable staging tables with full historical retention.\n\n"
              "Transformation layer: dbt models building SCD Type 2 dimensional tables—campaigns, "
              "accounts, assets, contacts—with proper slowly changing dimension handling so we could "
              "answer 'what did this campaign look like at the end of Q2?' without destroying history.\n\n"
              "Semantic layer: A Power BI dataset with pre-built DAX measures for the 25 most common "
              "metrics—campaign ROI, asset utilization rate, lead-to-opportunity conversion, budget "
              "burn rate. RLS (Row-Level Security) ensured account managers only saw their own accounts.\n\n"
              "The key decision was defining metrics in the semantic layer rather than in individual "
              "reports. That way, 'Campaign ROI' meant the same thing whether a VP or an analyst was "
              "looking at it.")

    add_label(doc, "Q:",
              "How did you drive adoption from 3 analyst users to 150+? What was your change "
              "management approach?")

    add_label(doc, "A:",
              "The technology was the easy part. The real work was behavioral change.\n\n"
              "I started by identifying power users—people who were already technically curious and "
              "frustrated with the status quo. I gave them early access and positioned them as internal "
              "champions. When their peers saw them answering questions in meetings by drilling into a "
              "dashboard live, instead of saying 'I'll get back to you,' it created organic demand.\n\n"
              "I also created 'Office Hours' sessions every other Friday—open Zoom calls where anyone "
              "could bring a data question and I'd show them how to answer it themselves in Power BI. "
              "Instead of giving them a fish, I taught them to filter a chart.\n\n"
              "For the skeptics, I attacked the trust problem directly. I built a data quality scorecard "
              "visible in the main dashboard—showing data freshness timestamps, row count reconciliation "
              "vs. source systems, and flagging any metric where something looked anomalous. When people "
              "could see the data was verified and up-to-date, they trusted it. The 60% reduction in "
              "IT data request tickets happened within 6 months of full deployment.")

    add_label(doc, "Q:", "Tell me about the campaign attribution modeling you built.")

    add_label(doc, "A:",
              "This was one of the more analytically interesting problems on the project. We had "
              "multi-touch campaign journeys—a prospect might see a display ad, receive an email, "
              "then get a call from a sales rep before converting. The business was using last-touch "
              "attribution, which gave 100% credit to the sales call and zero to the upstream digital "
              "touches.\n\n"
              "I built a data-driven attribution model in Python using a Markov chain approach—modeling "
              "customer journey sequences and calculating the removal effect of each channel (if you "
              "remove email from the path, how much does conversion probability drop?). The Markov "
              "attribution gave email campaigns 3x higher credit than last-touch, and display campaigns "
              "2x higher.\n\n"
              "This had a direct budget implication: the marketing team had been systematically "
              "underfunding email based on bad attribution data. After the model was deployed, they "
              "reallocated $400K toward email, which subsequently outperformed the prior campaign mix "
              "by 22% on a cost-per-lead basis.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # PROJECT 4 – Signet Jewelers Data Deduplication
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "4. ENTERPRISE DATA DEDUPLICATION PIPELINE – SIGNET JEWELERS", level=1,
                color=(30, 100, 170))
    p = doc.add_paragraph()
    p.add_run("150M+ Records  |  Zero Data Loss  |  AWS Glue / Redshift / Salesforce Bulk API  |  June–September 2025").bold = True

    add_label(doc, "Q:",
              "150 million records is a massive scale for a deduplication project. Set the scene—"
              "what was the data quality problem Signet hired you to solve?")

    add_label(doc, "A:",
              "Signet Jewelers had accumulated significant customer data debt from years of acquisitions—"
              "Zales, Kay Jewelers, Jared, all operating with their own legacy CRM instances before being "
              "integrated. When those systems merged into a central Salesforce CRM, the consolidation "
              "happened without a proper deduplication strategy. The result was 150M+ customer and "
              "transaction records where a single customer might exist as 4 or 5 separate records—"
              "different name spellings, multiple email addresses, address variations from moves over "
              "the years.\n\n"
              "The business impact was significant: marketing was sending duplicate campaign mailers, "
              "loyalty programs were fragmented so customers couldn't see their full purchase history, "
              "and the 360-degree customer view that personalization depends on was impossible. "
              "My contract was to design and execute the deduplication pipeline end-to-end.")

    add_label(doc, "Q:",
              "Walk me through the technical architecture. How do you even approach "
              "deduplication at 150 million records?")

    add_label(doc, "A:",
              "Scale forces you to be disciplined about your matching strategy. You cannot do an "
              "O(n²) pairwise comparison across 150M records—that's 22.5 trillion comparisons. "
              "You need blocking.\n\n"
              "The architecture used a three-stage approach in AWS:\n\n"
              "Stage 1 – Blocking: I used PySpark on AWS Glue to partition records into candidate "
              "blocks—groups of records that are plausibly the same entity. Blocking keys included "
              "Soundex-normalized last name + ZIP code, email domain + first 3 characters of last name, "
              "and phone number (normalized to E.164 format). This reduced the comparison space by "
              "~99.7% while retaining virtually all true duplicates.\n\n"
              "Stage 2 – Matching: Within each block, I computed feature-level similarity scores: "
              "Jaro-Winkler for name fields, exact match for email and phone, geolocation-adjusted "
              "address matching for addresses. A weighted composite score determined match probability. "
              "Deterministic rules (exact email OR exact phone + matching ZIP) auto-confirmed matches "
              "above the threshold. Probabilistic matches between 0.70 and 0.95 went to a human "
              "review sample.\n\n"
              "Stage 3 – Golden Record Creation: For confirmed duplicates, I built a survivorship "
              "rule engine in SQL—defining which source record's values 'survive' for each field. "
              "Most recent non-null email, most complete address, earliest creation date as customer "
              "since date, union of all purchase history records. This created the golden customer record "
              "that synced back to Salesforce via Bulk API 2.0.")

    add_label(doc, "Q:", "How did you ensure zero data loss throughout the process?")

    add_label(doc, "A:",
              "Data loss in a deduplication project is the worst-case outcome—you're permanently "
              "destroying records, so the audit framework has to be airtight.\n\n"
              "I built a multi-checkpoint reconciliation framework:\n\n"
              "Pre-processing baseline: Before any transformation, I captured exact row counts, "
              "null rates, value distributions, and referential integrity statistics for every "
              "table. These became the benchmark for all downstream validation.\n\n"
              "Post-blocking validation: After blocking, I verified that the total number of "
              "unique record IDs in all blocks equaled the source table count. No record should "
              "be silently dropped during blocking.\n\n"
              "Post-merge validation: After golden record creation, I ran a reconciliation query "
              "confirming that every source record ID was either: (a) promoted as a golden record, "
              "or (b) linked as a subordinate to exactly one golden record. No orphans allowed.\n\n"
              "Salesforce sync validation: After Bulk API 2.0 upserts, I compared Salesforce record "
              "counts vs. golden record counts. AWS CloudWatch alarms triggered if any batch had a "
              "failure rate above 0.01%.\n\n"
              "The reconciliation framework ran end-to-end and produced a signed-off data quality "
              "report before each production batch committed. Zero data loss through the entire "
              "150M record processing run.")

    add_label(doc, "Q:",
              "What were the trickiest edge cases in the matching logic, and how did "
              "you handle them?")

    add_label(doc, "A:",
              "Three edge cases stand out:\n\n"
              "Married name changes: A customer who bought an engagement ring as Jane Smith and "
              "returned for anniversary jewelry as Jane Johnson looked like two different people "
              "to a name-based matcher. I handled this by cross-referencing address and phone "
              "number continuity—if address and phone matched exactly, I accepted the match "
              "regardless of last name change, and flagged the record for manual review to confirm.\n\n"
              "Family members at the same address: A household might have three Johnsons at the "
              "same address—a father, mother, and adult child. Address-based blocking would throw "
              "all three into the same block, and naive matching would merge them. I added a "
              "disambiguation step: if first names differed AND email addresses differed AND phone "
              "numbers differed, they were treated as distinct customers even if address matched.\n\n"
              "Business vs. consumer accounts: Some jewelers sold to corporate accounts (companies "
              "buying gift cards or recognition jewelry). Business accounts would match consumer "
              "accounts if they shared a contact name. I added an account_type field to the survivorship "
              "rules—business accounts were never merged with consumer accounts regardless of name match.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # PROJECT 5 – Syilum Product Analytics
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "5. PRODUCT PERFORMANCE & USER ANALYTICS – SYILUM LLC", level=1,
                color=(30, 100, 170))
    p = doc.add_paragraph()
    p.add_run("8+ Years  |  60% Reduction in Ad-Hoc Requests  |  99%+ Report Accuracy  |  January 2016 – March 2024").bold = True

    add_label(doc, "Q:",
              "Eight years with the same company as a product analyst is a long tenure. "
              "What kept the work interesting, and how did your analytics function evolve over that time?")

    add_label(doc, "A:",
              "The product itself evolved through multiple platform versions, so the analytical "
              "questions kept pace with the product's maturity.\n\n"
              "In the early years, the core questions were foundational: who are our users, what "
              "features are they using, where do they drop off? I was building the analytics "
              "infrastructure from scratch—writing SQL against the raw application database, "
              "building the first Power BI workspace, establishing the KPI taxonomy.\n\n"
              "By the middle years, we had enough historical data to do longitudinal analysis—cohort "
              "retention curves, feature adoption trajectories, user segment behavioral clustering. "
              "I introduced IBM Cloud Pak for Data (CP4D) with Watson Query to enable distributed "
              "analytics across our multi-environment infrastructure.\n\n"
              "In the final years, my role shifted from doing the analysis to enabling the organization "
              "to do its own analysis. I built the self-service analytics layer in Power BI—a semantic "
              "model with guided exploration paths that reduced ad-hoc requests by 60% and freed me "
              "to focus on strategic analysis instead of repetitive SQL queries.")

    add_label(doc, "Q:",
              "Tell me about your A/B testing practice. How did you design and analyze "
              "experiments at Syilum?")

    add_label(doc, "A:",
              "A/B testing at a software company with a relatively small user base requires more "
              "statistical rigor than most people expect—you can't be sloppy about power calculations "
              "or you'll run underpowered tests and draw false conclusions.\n\n"
              "My standard process:\n\n"
              "Pre-experiment: I worked with product managers to define the primary metric (what we're "
              "trying to move), the minimum detectable effect (what change is actually meaningful), "
              "and the expected baseline conversion rate. I'd then calculate the required sample size "
              "using power analysis (typically 80% power, α=0.05) to determine how long the test needed "
              "to run.\n\n"
              "During the experiment: I built daily monitoring dashboards showing sample accumulation "
              "per variant, guardrail metrics (we never wanted to improve one metric while secretly "
              "harming another), and a novelty effect check—watching for the initial engagement spike "
              "that appears with any UI change and settles after 3–4 days.\n\n"
              "Post-experiment: I used a two-proportion z-test for conversion rate experiments and "
              "a t-test for continuous metrics. I also computed practical significance alongside "
              "statistical significance—a 0.2% conversion improvement might be statistically significant "
              "but not worth the engineering effort to maintain.\n\n"
              "One standout result: an A/B test on the onboarding wizard revealed that adding a "
              "single tooltip to the Advanced Analytics setup step increased feature adoption from "
              "22% to 38% within 30 days. A $0 fix with the most measurable product impact of the year.")

    add_label(doc, "Q:",
              "You mentioned discovering that users who engaged with 3+ features in their first week "
              "had 4x higher 90-day retention. How did you find that insight and what happened with it?")

    add_label(doc, "A:",
              "It was a pattern I noticed while building the feature adoption dashboard—not something "
              "anyone had asked me to find. I had a cohort table in Snowflake showing week-1 feature "
              "engagement counts per user. I was running a standard retention analysis alongside it "
              "and noticed the retention curves were dramatically different based on week-1 engagement "
              "breadth.\n\n"
              "I added a calculated column: features_used_week_1, bucketed into 1, 2, and 3+ categories. "
              "When I plotted 90-day retention against those buckets, the separation was stark—1-feature "
              "users retained at 34%, 2-feature at 58%, and 3+ feature users at 71%. The 4x headline "
              "compared the bottom and top buckets.\n\n"
              "I shared this in a product review with no fanfare—just a chart and a question: 'Can we "
              "identify users in their first week who haven't hit 3 features yet and nudge them?' The VP "
              "of Product immediately saw the implications. Within 6 weeks, engineering shipped an "
              "in-app nudge campaign for new users below the 3-feature threshold. That report became "
              "the most-viewed in the entire Power BI workspace.\n\n"
              "The lesson I took from that: the most impactful analysis is often the question nobody "
              "thought to ask, surfaced from exploratory work that isn't in any project plan.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # PROJECT 6 – AT&T / YP Marketing Campaign Optimization
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "6. MARKETING CAMPAIGN ANALYTICS – AT&T / YP ADVERTISING SOLUTIONS",
                level=1, color=(30, 100, 170))
    p = doc.add_paragraph()
    p.add_run("20% ROAS Improvement  |  100+ Campaigns Analyzed  |  +12% Renewal Conversion  |  January 2010 – December 2013").bold = True

    add_label(doc, "Q:",
              "This was an earlier role in your career at YP. What does a marketing data analyst "
              "do in a local advertising business, and what analytical problems were you solving?")

    add_label(doc, "A:",
              "YP (Yellow Pages) at that time was in the middle of a pivotal transition—print "
              "directories were declining and digital advertising was ascendant, but we still had "
              "8,200+ active advertiser accounts across 12 markets who were buying both print and "
              "digital. My job was to make sense of whether those advertising dollars were actually "
              "working for advertisers, and to give the sales and renewal teams data-backed tools "
              "to have better conversations.\n\n"
              "The core problems were:\n")
    add_bullet_list(doc, [
        "No unified view of campaign performance across print and digital channels—each lived in separate systems",
        "Weekly campaign reports took 2–3 days to compile manually in Excel and were frequently wrong",
        "The renewal team had no predictive signal for which advertisers were about to cancel—they operated on gut feel",
        "Inconsistent advertiser IDs across billing and CRM caused revenue mismatches that delayed month-end close",
    ])

    add_label(doc, "Q:",
              "Tell me about the multi-channel attribution model you built. "
              "How did you reconcile print and digital attribution?")

    add_label(doc, "A:",
              "This is the analysis I'm most proud of from that period, and it taught me a "
              "lesson I still use today: attribution only works when you measure all the channels "
              "that matter.\n\n"
              "When I first built the attribution model, digital looked overwhelmingly superior "
              "to print on a last-touch basis—search ads converted, print didn't seem to. "
              "The sales team was gearing up to aggressively push advertisers off print and onto "
              "digital-only packages. But the data was lying through omission.\n\n"
              "Print ads don't generate clicks. They generate calls. And call volume wasn't being "
              "fed back into the attribution model at all. I integrated call tracking data from "
              "our telephony vendor and built a phone-call attribution layer—matching inbound calls "
              "to the print ads that ran in the corresponding market and time window.\n\n"
              "When call attribution was included, print was driving 31% of verifiable customer "
              "contacts for local service businesses—plumbers, electricians, restaurants—who didn't "
              "track online conversions at all. The 'print is dying' narrative in the data was a "
              "measurement problem, not a market reality. That analysis directly influenced the "
              "product team's packaging strategy and prevented a premature print deprecation that "
              "would have hurt advertiser retention.")

    add_label(doc, "Q:",
              "How did you build the advertiser churn propensity model, and how did "
              "it improve renewal conversion by 12%?")

    add_label(doc, "A:",
              "The renewal team was managing 8,200+ accounts with no systematic way to prioritize "
              "outreach. Good account managers had intuition, but intuition doesn't scale.\n\n"
              "I built an RFM (Recency, Frequency, Monetary) segmentation first as a baseline—"
              "scoring each advertiser on how recently they'd had a performance review, how many "
              "active campaigns they had, and their monthly spend level. High-RFM accounts were "
              "safe; low-RFM were at risk.\n\n"
              "I then built a churn propensity regression model layered on top of RFM, adding "
              "features like: average CTR trend over trailing 90 days (declining CTR predicts churn), "
              "call attribution volume (zero calls = zero perceived value = churn risk), time since "
              "last account manager contact, and whether the advertiser had raised a billing dispute "
              "in the prior 6 months (strong churn signal).\n\n"
              "The model scored every account 60 days before renewal and assigned a risk tier: High, "
              "Medium, Low. Renewal reps received a prioritized outreach list with data-backed talking "
              "points for each at-risk account—'Your search ads CTR dropped 40% last quarter; let's "
              "talk about creative refresh.' That specificity changed the quality of renewal conversations "
              "from generic check-ins to consultative problem-solving.\n\n"
              "Renewal conversion improved 12% in the two quarters following deployment.")

    add_label(doc, "Q:",
              "You mentioned establishing a data dictionary for 47 business metrics. "
              "Why was that necessary and how did you execute it?")

    add_label(doc, "A:",
              "It became necessary because of a humiliating meeting. The VP of Marketing asked me "
              "to pull 'total active advertisers.' I returned a number. Finance returned a different "
              "number. Operations had a third number. Three different answers to the same question—"
              "because 'active' meant something different in every system.\n\n"
              "In billing, active meant 'has an open invoice.' In CRM, it meant 'account is not "
              "marked cancelled.' In fulfillment, it meant 'at least one ad is currently running.' "
              "All three were defensible. None were the same.\n\n"
              "The data dictionary project came out of that meeting. But I learned a critical lesson "
              "about governance execution: you can build the best data dictionary in the world and it "
              "means nothing if people don't use it. My first version was a beautifully formatted "
              "Excel file. Nobody opened it.\n\n"
              "The turning point was when I stopped documenting and started facilitating. Instead of "
              "defining 'active advertiser' for everyone, I brought finance, sales ops, and fulfillment "
              "into a room and made them argue it out. When they reached consensus, I wrote it down. "
              "Suddenly it was 'their' definition, not mine. Adoption happened organically because the "
              "people who would use the data had ownership of the definition.\n\n"
              "Governance is change management, not data management. That insight has guided every "
              "governance initiative I've led since.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # PROJECT 7 – Support Optics Healthcare Analytics
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "7. HEALTHCARE DATA ANALYTICS – SUPPORT OPTICS (CONTRACT)",
                level=1, color=(30, 100, 170))
    p = doc.add_paragraph()
    p.add_run("HIPAA Compliant  |  Predictive Patient Outcome Modeling  |  Multi-Sector (Finance, BI, Healthcare)  |  January 2013 – January 2016").bold = True

    add_label(doc, "Q:",
              "Healthcare analytics has unique constraints around data privacy. How did you approach "
              "working with sensitive patient data while still building meaningful predictive models?")

    add_label(doc, "A:",
              "HIPAA compliance wasn't a checkbox—it was a design constraint that shaped every technical "
              "decision from storage architecture to query patterns.\n\n"
              "The data access model used role-based access control at the column level. PHI fields "
              "(patient name, date of birth, SSN, address) were stored in a separate encrypted schema "
              "with a different access tier than the analytical fields (diagnosis codes, procedure codes, "
              "outcome flags, admission/discharge timestamps). Analysts working on quality improvement "
              "models accessed de-identified datasets with a token-based patient ID. The PHI linkage "
              "table was only accessible to authorized data stewards for specific, logged purposes.\n\n"
              "For model development specifically, I worked exclusively with de-identified data using "
              "Safe Harbor de-identification (removing 18 identifiers per the HIPAA Safe Harbor standard). "
              "This meant I could develop and validate predictive models without ever touching PHI. "
              "When predictions needed to be operationalized—for example, surfacing high-risk patients "
              "to care coordinators—the final join back to patient identifiers happened inside the "
              "secure environment, not in the model scoring pipeline.")

    add_label(doc, "Q:",
              "Tell me about the patient outcome predictive models you built. "
              "What were you trying to predict and what was the clinical significance?")

    add_label(doc, "A:",
              "The primary model I built was a readmission risk classifier—predicting which patients "
              "were at elevated risk of being readmitted within 30 days of hospital discharge. "
              "This is a well-established quality metric in healthcare; high 30-day readmission rates "
              "are both costly (CMS penalizes hospitals for excessive readmissions under the HRRP program) "
              "and clinically meaningful (readmissions often indicate inadequate discharge planning or "
              "follow-up care).\n\n"
              "Features I engineered from the available data:\n")
    add_bullet_list(doc, [
        "Clinical complexity: number of active diagnoses (Elixhauser comorbidity index), prior admission count in trailing 12 months",
        "Discharge context: discharge disposition (home, skilled nursing, rehab), length of stay relative to DRG average",
        "Social determinants: insurance type (proxy for socioeconomic status), distance from facility",
        "Procedure mix: whether the admission involved a surgical procedure, emergency vs. elective admission flag",
        "Lab value trends: trending deterioration in key labs (creatinine, sodium) in the 48 hours before discharge",
    ])
    p = doc.add_paragraph(
        "I used a Random Forest classifier with SMOTE oversampling to handle class imbalance "
        "(readmissions were ~15% of the population). The model achieved 78% AUC-ROC on the validation "
        "set, which compared favorably against the LACE index (a standard clinical risk tool) at ~72% "
        "AUC on the same dataset. High-risk flags were surfaced to discharge planning nurses 24 hours "
        "before discharge to trigger enhanced transition-of-care protocols."
    )
    p.runs[0].font.size = Pt(11)

    add_label(doc, "Q:",
              "How did you handle the statistical analysis and hypothesis testing "
              "in a healthcare quality improvement context?")

    add_label(doc, "A:",
              "Healthcare quality improvement requires more methodological caution than typical "
              "business analytics because poor analysis can lead to clinical decisions.\n\n"
              "I used regression analysis for the primary outcome studies—logistic regression for "
              "binary outcomes (readmission yes/no), linear regression for continuous outcomes "
              "(length of stay, cost per episode). For comparative analysis—comparing outcome rates "
              "between patient cohorts or time periods—I used chi-square tests for categorical outcomes "
              "and t-tests or Mann-Whitney U for continuous outcomes depending on distributional assumptions.\n\n"
              "One area I was particularly careful about was confounding. If we saw lower readmission "
              "rates after implementing a new discharge protocol, we couldn't simply attribute that "
              "to the protocol—patient mix changes over time, seasonal effects exist, and the care "
              "teams implementing the protocol were also receiving training simultaneously. I used "
              "difference-in-differences analysis where we had a comparable control population, and "
              "I was explicit in reporting about the observational (non-experimental) nature of the "
              "analysis and the limitations that imposed on causal inference.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # PROJECT 8 – Exoplanet Habitability (Personal Project)
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "8. EXOPLANET HABITABILITY ANALYSIS – PERSONAL PROJECT",
                level=1, color=(30, 100, 170))
    p = doc.add_paragraph()
    p.add_run("5,000+ Exoplanets Analyzed  |  NASA Kepler Data  |  Multi-Algorithm ML Classification  |  November 2025").bold = True

    add_label(doc, "Q:",
              "Tell me about this personal project. Exoplanet habitability feels like a departure "
              "from business analytics—why did you pursue it and what were you trying to learn?")

    add_label(doc, "A:",
              "I built this project for two reasons: intellectual curiosity and technical skill "
              "development. The astronomy aspect has always fascinated me—there's something compelling "
              "about using the same ML techniques I apply to campaign optimization to ask whether "
              "there are Earth-like worlds around other stars.\n\n"
              "On the technical side, I wanted to practice the full ML pipeline in a context where I "
              "owned every decision—from problem framing through feature engineering to model deployment "
              "and interpretation—without the organizational constraints that sometimes limit what you "
              "can explore in a professional setting. It's also a portfolio piece that demonstrates I "
              "can work with complex scientific data and ambiguous problem definitions, not just "
              "well-structured business datasets.\n\n"
              "The core challenge was definitional: what does 'habitable' mean? I anchored to the "
              "circumstellar habitable zone (the 'Goldilocks zone' where liquid water can exist on "
              "a planet surface) as my primary proxy, combined with radius constraints filtering for "
              "terrestrial-sized planets. That gave me a defensible binary target variable to build "
              "a classification model around.")

    add_label(doc, "Q:",
              "Walk me through your ML approach—what algorithms did you use and "
              "how did you handle the feature engineering for astronomical data?")

    add_label(doc, "A:",
              "I worked with 5,000+ confirmed exoplanets from the NASA Kepler dataset. The raw "
              "features were astronomical measurements:\n")
    add_bullet_list(doc, [
        "Orbital period (days)",
        "Planet radius (Earth radii)",
        "Planet mass (Earth masses, where available—significant missing data)",
        "Semi-major axis (AU)",
        "Orbital eccentricity",
        "Host stellar mass (solar masses)",
        "Host stellar effective temperature (Kelvin)",
        "Distance from Earth (parsecs)",
    ])
    p = doc.add_paragraph(
        "Feature engineering added physically meaningful derived features: equilibrium temperature "
        "(calculated from stellar luminosity and semi-major axis), the 'Earth Similarity Index' for "
        "radius and stellar temperature, and a habitable zone flag based on the Kopparapu et al. "
        "boundaries for the host star's luminosity class.\n\n"
        "Planet mass had ~40% missingness because the radial velocity measurements needed to calculate "
        "mass aren't available for all Kepler candidates. I used a random forest imputer (iterative "
        "imputation) to predict missing mass values from the available radius, density proxy, and "
        "stellar characteristics.\n\n"
        "I trained four models: Logistic Regression (baseline), Random Forest, Gradient Boosting "
        "(XGBoost), and a neural network (MLPClassifier with 3 hidden layers). Given the severe "
        "class imbalance—only ~2% of exoplanets are classified as potentially habitable—I used "
        "SMOTE on the training set and evaluated on imbalanced holdout data using F1 (weighted) "
        "and AUC-ROC as primary metrics rather than accuracy. The Gradient Boosting model performed "
        "best with an AUC-ROC of 0.94 and F1 of 0.89 on the holdout set."
    )
    p.runs[0].font.size = Pt(11)

    add_label(doc, "Q:",
              "What did the model tell you about which features most predict habitability? "
              "And did any findings surprise you?")

    add_label(doc, "A:",
              "SHAP analysis identified the top predictors as: equilibrium temperature (unsurprising—"
              "the habitable zone is defined by temperature), planet radius (terrestrial planets are "
              "small—above ~1.5 Earth radii you're likely looking at a mini-Neptune with a thick "
              "hydrogen envelope), and host stellar effective temperature (K-type and G-type stars "
              "are more habitable than M-dwarfs or hot A-type stars).\n\n"
              "The finding that surprised me most was the importance of eccentricity. High orbital "
              "eccentricity means a planet swings from very close to very far from its star during "
              "each orbit—creating extreme seasonal temperature variation. I expected this to be a "
              "moderate factor. SHAP values showed it was the fourth most important feature, stronger "
              "than stellar mass. Planets in the habitable zone with eccentricity above 0.3 had "
              "dramatically lower predicted habitability, even if their average distance placed them "
              "in the zone. That's physically intuitive when you think about it—time-averaged "
              "temperature isn't the same as stable-within-range temperature—but I hadn't expected "
              "the magnitude of the effect to show up so clearly in the data.\n\n"
              "The galactic extrapolation was also memorable: applying the model's estimated 2.1% "
              "habitability rate to the ~400 billion stars in the Milky Way suggested roughly "
              "8.4 billion potentially habitable Earth-like planets in our galaxy alone. Even "
              "accounting for enormous model uncertainty, that number is staggering.")

    add_label(doc, "Q:",
              "How do you connect a personal passion project like this to your professional work? "
              "Why should an interviewer care about it?")

    add_label(doc, "A:",
              "The skills transfer directly, even if the domain doesn't. Any time you have an "
              "ambiguous target variable, severe class imbalance, significant missing data, and "
              "the need to explain complex model behavior to a non-technical audience—you're facing "
              "the same challenges whether the dataset is exoplanets or customer churn.\n\n"
              "Specifically, the missing data imputation strategy I developed for planet mass became "
              "a template I used at AT&T for handling missing campaign budget data. The SHAP "
              "visualization approach I refined for explaining habitability features to a non-specialist "
              "audience directly improved how I presented ML model explanations to marketing executives.\n\n"
              "More broadly, I build personal projects because I believe data science is a craft, "
              "and craft improves with deliberate practice outside of work constraints. Someone who "
              "only does data science when assigned to is not as curious as someone who does it "
              "because they can't help it.")

    doc.add_page_break()

    # ════════════════════════════════════════════════════════════════════════
    # GENERAL / CROSS-CUTTING QUESTIONS
    # ════════════════════════════════════════════════════════════════════════
    add_heading(doc, "9. GENERAL INTERVIEW QUESTIONS – CROSS-PROJECT", level=1,
                color=(30, 100, 170))
    doc.add_paragraph("These are follow-up questions interviewers ask across all projects.")

    add_label(doc, "Q:",
              "You've worked across healthcare, telecom, retail, and product analytics. "
              "How do you ramp up quickly in a new domain?")

    add_label(doc, "A:",
              "My approach has three phases.\n\n"
              "First, I spend the first week listening, not building. I schedule 30-minute "
              "conversations with people across the business—product, finance, operations, sales—"
              "and ask the same two questions: 'What question do you ask most that you can never "
              "get a good answer to?' and 'What do you wish the data team understood better about "
              "your work?' That surfaces the real problems faster than any requirements document.\n\n"
              "Second, I get my hands on the data immediately and generate a data quality report "
              "before doing any analysis. Understanding the shape of the data—completeness, "
              "distributions, referential integrity—prevents me from building models on faulty "
              "foundations.\n\n"
              "Third, I identify one quick win to deliver in the first 30 days. Not because quick "
              "wins are the goal, but because they build the stakeholder trust that earns you "
              "permission to tackle the harder, longer-term problems.")

    add_label(doc, "Q:",
              "What's your philosophy on choosing between a simple model and a complex one?")

    add_label(doc, "A:",
              "I start with the simplest model that could plausibly work and only add complexity "
              "when I can quantify the improvement it provides.\n\n"
              "In practice that means: logistic regression before Random Forest, Random Forest "
              "before XGBoost, XGBoost before neural networks. Each step up in complexity costs "
              "you interpretability, maintainability, and the ability to explain predictions to "
              "stakeholders. That cost has to be justified by measurable performance gain.\n\n"
              "There's also a deployment consideration. A logistic regression model can be "
              "implemented as a SQL query in a data warehouse—no ML infrastructure required, "
              "any analyst can maintain it. An XGBoost model requires a serving endpoint, "
              "version management, and specialized skills to retrain. If the logistic regression "
              "gives you 82% and XGBoost gives you 84%, the infrastructure cost difference "
              "usually doesn't justify the uplift.\n\n"
              "I've also found that the 'explainability conversation' with business stakeholders "
              "is as important as model performance. An 89% accurate black box that stakeholders "
              "don't trust will not change behavior. An 82% accurate model where you can say "
              "'the top 3 drivers of this prediction are X, Y, Z' gets adopted and drives action.")

    add_label(doc, "Q:",
              "How do you communicate ML results to non-technical executives? "
              "Walk me through your actual process.")

    add_label(doc, "A:",
              "I never lead with the model. I lead with the business outcome.\n\n"
              "My standard executive presentation structure:\n\n"
              "1. Problem reminder (1 slide): Here's the business pain we were solving—in their "
              "language, with their numbers.\n\n"
              "2. The answer (1 slide): Here's what we found. The model says X% of campaigns are "
              "at risk. Here are the three you should focus on today.\n\n"
              "3. Why we believe it (2 slides): Here's how the model was validated (I translate "
              "'AUC-ROC' to 'out of every 10 campaigns that actually overran, the model flagged 8.5 "
              "of them in advance'). Here are the top 3 factors driving the prediction.\n\n"
              "4. Action (1 slide): Here's what I'm recommending we do differently as a result.\n\n"
              "I keep model details—accuracy metrics, algorithm choices, cross-validation methodology—"
              "in an appendix. Executives who want that depth will ask for it. Most don't. What they "
              "need is confidence that the answer is reliable and clarity on what to do with it.\n\n"
              "The metric translation practice is essential: nobody outside of data science knows "
              "what AUC-ROC is, but everyone understands 'we catch 85 out of every 100 at-risk "
              "campaigns before they overrun.'")

    add_label(doc, "Q:", "What's the most important lesson you've learned across 15+ years of data work?")

    add_label(doc, "A:",
              "Data quality is a people problem, not a technology problem.\n\n"
              "Every data quality issue I've ever encountered traces back to human decisions: a "
              "developer who didn't validate input, a process that evolved without updating the schema, "
              "two teams who defined the same metric differently because they never talked to each other.\n\n"
              "Technology is how you detect and remediate data quality problems. People are why they "
              "exist in the first place. If you build the most sophisticated monitoring pipeline in "
              "the world but don't change the upstream behaviors that create the problems, you're "
              "running to stand still.\n\n"
              "The most effective data governance work I've done has been facilitation—getting the "
              "finance team, the operations team, and the sales team in the same room to agree on "
              "what 'active customer' means, and then giving each person ownership of their piece "
              "of the definition. That meeting does more for data quality than any validation "
              "framework I could write.\n\n"
              "The technical skills matter. But the meta-skill—understanding that data is a "
              "sociotechnical system, not just a technical one—is what separates analysts who "
              "produce outputs from analysts who produce change.")

    # ── Footer ──────────────────────────────────────────────────────────────
    doc.add_page_break()
    footer_p = doc.add_paragraph(
        "Pamela Austin | Data Science & Analytics Portfolio | ML Project Interview Q&A\n"
        "Built from live portfolio projects: AT&T, Signet Jewelers, Syilum LLC, Support Optics, "
        "YP Advertising Solutions, and Personal Research."
    )
    footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in footer_p.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(130, 130, 130)
        run.italic = True

    # ── Save ─────────────────────────────────────────────────────────────────
    output_path = os.path.join(
        os.path.dirname(__file__),
        "Pamela_Austin_ML_Interview_QA.docx"
    )
    doc.save(output_path)
    print(f"Document saved to: {output_path}")
    return output_path


if __name__ == "__main__":
    build_document()
