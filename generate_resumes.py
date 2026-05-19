"""
Generate three professional Word resumes for Pamela Austin based on her portfolio.
Resumes: (1) Senior Data & BI Analyst, (2) Data Program Manager, (3) Analytics Engineer
"""
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))

# ── Color palette ────────────────────────────────────────────────────────────
DARK_BLUE   = RGBColor(0x1A, 0x37, 0x5E)   # header / section titles
MID_BLUE    = RGBColor(0x2E, 0x5F, 0x9B)   # sub-headers
ACCENT      = RGBColor(0x14, 0xB8, 0xA6)   # teal accent line
BODY_BLACK  = RGBColor(0x1A, 0x1A, 0x1A)
GRAY        = RGBColor(0x55, 0x65, 0x70)


# ── Low-level XML helpers ────────────────────────────────────────────────────

def add_bottom_border(paragraph, color="14B8A6", size=12):
    """Add a colored bottom border to a paragraph (used as a section divider)."""
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(size))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def set_cell_bg(cell, color_hex):
    """Set table cell background shading."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def remove_table_borders(table):
    """Remove all borders from a table."""
    tbl = table._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement("w:tblBorders")
    for border_name in ("top", "left", "bottom", "right", "insideH", "insideV"):
        b = OxmlElement(f"w:{border_name}")
        b.set(qn("w:val"), "none")
        tblBorders.append(b)
    tblPr.append(tblBorders)


# ── Document factory ─────────────────────────────────────────────────────────

def make_document():
    doc = Document()
    # Narrow margins (0.75 in all around)
    for section in doc.sections:
        section.top_margin    = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin   = Inches(0.75)
        section.right_margin  = Inches(0.75)
    # Remove default Normal paragraph spacing
    style = doc.styles["Normal"]
    style.font.name  = "Calibri"
    style.font.size  = Pt(10)
    style.font.color.rgb = BODY_BLACK
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after  = Pt(0)
    return doc


# ── Shared content builders ──────────────────────────────────────────────────

def add_name_header(doc, name, title, contact_line):
    """Big name, title subtitle, then contact on one line."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(name)
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Calibri"

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_after = Pt(2)
    r2 = p2.add_run(title)
    r2.font.size = Pt(12)
    r2.font.color.rgb = MID_BLUE
    r2.font.name = "Calibri"
    r2.bold = True

    p3 = doc.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_after = Pt(6)
    r3 = p3.add_run(contact_line)
    r3.font.size = Pt(9)
    r3.font.color.rgb = GRAY
    r3.font.name = "Calibri"

    # Teal divider
    div = doc.add_paragraph()
    div.paragraph_format.space_before = Pt(0)
    div.paragraph_format.space_after  = Pt(8)
    add_bottom_border(div, "1A375E", 18)


def add_section_heading(doc, heading_text, space_before=10):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after  = Pt(3)
    run = p.add_run(heading_text.upper())
    run.bold = True
    run.font.size = Pt(10)
    run.font.color.rgb = DARK_BLUE
    run.font.name = "Calibri"
    add_bottom_border(p, "14B8A6", 8)
    return p


def add_job(doc, company, role, dates, bullets):
    """Add one job block: company/role/dates, then bullet list."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(1)

    r_company = p.add_run(company)
    r_company.bold = True
    r_company.font.size = Pt(10.5)
    r_company.font.color.rgb = BODY_BLACK
    r_company.font.name = "Calibri"

    r_sep = p.add_run("  |  ")
    r_sep.font.size = Pt(10)
    r_sep.font.color.rgb = GRAY

    r_role = p.add_run(role)
    r_role.font.size = Pt(10)
    r_role.font.color.rgb = MID_BLUE
    r_role.bold = True

    # Dates right-aligned via tab stop
    r_tab = p.add_run("\t")
    r_dates = p.add_run(dates)
    r_dates.font.size = Pt(9.5)
    r_dates.italic = True
    r_dates.font.color.rgb = GRAY

    # Set a right-aligned tab stop at ~7 inches from left margin
    pPr = p._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), "9360")   # 9360 twips = 6.5 in (page width minus margins)
    tabs.append(tab)
    pPr.append(tabs)

    for bullet in bullets:
        bp = doc.add_paragraph(style="List Bullet")
        bp.paragraph_format.left_indent   = Inches(0.2)
        bp.paragraph_format.space_before  = Pt(1)
        bp.paragraph_format.space_after   = Pt(1)
        br = bp.add_run(bullet)
        br.font.size = Pt(9.5)
        br.font.color.rgb = BODY_BLACK
        br.font.name = "Calibri"


def add_skills_table(doc, skills_dict):
    """Two-column skills table with category labels."""
    items = list(skills_dict.items())
    # Pair up categories for a two-column layout
    rows_needed = (len(items) + 1) // 2
    table = doc.add_table(rows=rows_needed, cols=4)
    table.style = "Table Grid"
    remove_table_borders(table)
    # Set column widths
    for row in table.rows:
        row.cells[0].width = Inches(1.1)
        row.cells[1].width = Inches(2.6)
        row.cells[2].width = Inches(1.1)
        row.cells[3].width = Inches(2.6)

    for idx, (category, skills_text) in enumerate(items):
        row_idx = idx // 2
        col_offset = (idx % 2) * 2
        cell_label = table.cell(row_idx, col_offset)
        cell_value = table.cell(row_idx, col_offset + 1)

        p_label = cell_label.paragraphs[0]
        p_label.paragraph_format.space_before = Pt(2)
        p_label.paragraph_format.space_after  = Pt(2)
        rl = p_label.add_run(category)
        rl.bold = True
        rl.font.size = Pt(9)
        rl.font.color.rgb = MID_BLUE
        rl.font.name = "Calibri"

        p_val = cell_value.paragraphs[0]
        p_val.paragraph_format.space_before = Pt(2)
        p_val.paragraph_format.space_after  = Pt(2)
        rv = p_val.add_run(skills_text)
        rv.font.size = Pt(9)
        rv.font.color.rgb = BODY_BLACK
        rv.font.name = "Calibri"


def add_education(doc, entries):
    for deg, school, year in entries:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after  = Pt(1)
        r1 = p.add_run(deg)
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.color.rgb = BODY_BLACK
        r2 = p.add_run(f"  –  {school}")
        r2.font.size = Pt(10)
        r2.font.color.rgb = GRAY
        r3 = p.add_run(f"  ({year})")
        r3.italic = True
        r3.font.size = Pt(9.5)
        r3.font.color.rgb = GRAY


def add_summary(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(6)
    r = p.add_run(text)
    r.font.size = Pt(10)
    r.font.color.rgb = BODY_BLACK
    r.font.name = "Calibri"


CONTACT = "pamtekk@gmail.com  |  linkedin.com/in/pamela-austin-621a32a4  |  pamelaaustin.dev"

EDU = [
    ("B.S. Computer Science", "National American University", "2016"),
    ("A.S. Finance & Accounting", "University of Phoenix", "2011"),
    ("IBM Certificate", "Database and SQL for Data Science with Python", "2023"),
]


# ════════════════════════════════════════════════════════════════════════════
# RESUME 1 — Senior Data Analyst / Business Intelligence Analyst
# ════════════════════════════════════════════════════════════════════════════

def build_resume_data_analyst():
    doc = make_document()

    add_name_header(
        doc,
        "Pamela Austin",
        "Senior Data Analyst  |  Business Intelligence Analyst",
        CONTACT,
    )

    # Summary
    add_section_heading(doc, "Professional Summary", space_before=0)
    add_summary(doc,
        "Results-driven Senior Data Analyst with 13+ years of experience translating complex datasets into "
        "actionable business intelligence across telecom, finance, retail, and technology. Expert in Power BI "
        "(DAX, RLS, drill-through), advanced SQL, dimensional data modeling, and self-service analytics "
        "platforms. Proven track record delivering measurable outcomes: $12.6M in accounting errors identified, "
        "92% faster financial consolidation, 93% dashboard query optimization, and ML forecasting models at "
        "94%+ accuracy. Trusted by C-suite and executive stakeholders to translate complex analytical findings "
        "into clear, actionable narratives."
    )

    # Skills
    add_section_heading(doc, "Core Competencies")
    add_skills_table(doc, {
        "BI & Visualization":   "Power BI (DAX, RLS, Drill-Through), Tableau, Executive Dashboards",
        "Databases & Cloud":    "Snowflake, SQL Server, Databricks, AWS, Fivetran, dbt",
        "Languages":            "SQL (Advanced), Python (Pandas, NumPy, scikit-learn), PySpark",
        "Data Modeling":        "Star Schema, Dimensional Modeling, Fact/Dimension Tables",
        "Analytics":            "Predictive Modeling, Forecasting, A/B Testing, Attribution Modeling",
        "Data Governance":      "Data Quality Frameworks, KPI Dictionaries, Stewardship Processes",
        "Integrations":         "Salesforce, Workfront, ZoomInfo, SAP",
        "Business Domains":     "Marketing Analytics, Financial Analytics, M&A, Product Analytics",
    })

    # Experience
    add_section_heading(doc, "Professional Experience")

    add_job(doc,
        "Syilum LLC", "Freelance Business Intelligence Analyst", "Oct 2025 – Present",
        [
            "Build Power BI dashboards tracking product KPIs (user engagement, feature adoption, retention) for "
            "product management and executive stakeholders, achieving 99%+ data accuracy.",
            "Write advanced SQL queries (Snowflake) analyzing user behavior patterns, conversion funnels, and "
            "feature usage to inform product roadmap priorities.",
            "Developed a governed KPI dictionary resolving 23 dashboards with conflicting metric definitions, "
            "establishing a single source of truth for executive reporting.",
            "Created self-service analytics solutions reducing ad-hoc report requests by 60%, freeing engineering "
            "and product teams for strategic work.",
            "Designed data quality validation frameworks ensuring consistent, audit-ready executive reporting.",
        ],
    )

    add_job(doc,
        "Signet Jewelers", "Data Analyst – M&A Financial BI", "Jun 2025 – Sep 2025",
        [
            "Delivered BI and SQL analytics for the Corporate Development team supporting a $450M+ M&A "
            "acquisition pipeline; sole BI analyst responsible for all financial dashboard delivery.",
            "Built Power BI dashboards with drill-through reports and row-level security (RLS) for 50+ "
            "stakeholders, visualizing EBITDA trends, working capital analysis, and acquisition scenarios.",
            "Designed star schema dimensional model (12 fact tables, 18 dimensions) and automated SQL "
            "reconciliation logic that identified $12.6M in accounting errors including a $3.2M price "
            "renegotiation on a single acquisition deal.",
            "Reduced financial consolidation from 6 weeks to 5 days (92% improvement) through automated "
            "SQL reporting pipelines processing 15.2M+ financial records.",
            "Improved dashboard query performance by 65% through optimized DAX measures and data model tuning.",
            "Translated complex financial analytics into C-suite narratives via Power BI and PowerPoint.",
        ],
    )

    add_job(doc,
        "AT&T", "Senior Business Intelligence Analyst", "Mar 2024 – Mar 2025",
        [
            "Built AT&T Marketing Operations BI platform: 25+ Power BI dashboards for executive leadership "
            "tracking campaign performance, workforce costs, and budget allocations across 200+ concurrent campaigns.",
            "Optimized dashboard query times from 45 seconds to 3 seconds (93% improvement) through "
            "data model optimization, DAX tuning, and Snowflake query refactoring.",
            "Created self-service analytics enabling 150+ business users, reducing IT report requests by 60% "
            "and manual data retrieval by 40%.",
            "Consolidated Workfront, Salesforce, and ZoomInfo data via Fivetran into a unified Snowflake "
            "reporting layer using dbt transformations.",
            "Developed predictive ML model (Python/scikit-learn) forecasting project timelines at 89% accuracy, "
            "reducing project overruns by 28% and leadership escalations by 30%.",
            "Built campaign attribution logic and budget variance dashboards tracking 200+ concurrent campaigns "
            "with daily refresh.",
        ],
    )

    add_job(doc,
        "Syilum LLC", "Product Data Analyst", "Jan 2016 – Mar 2024",
        [
            "Delivered product and marketing business intelligence for a technology company over 8 years, "
            "building 100+ dashboards and supporting 10M+ daily event analytics.",
            "Designed and maintained self-service analytics infrastructure reducing IT request volume by 60%.",
            "Led A/B testing programs driving 23% feature adoption improvements; analyzed product funnels and "
            "cohort retention to inform roadmap decisions.",
            "Built ETL pipelines and dimensional data models supporting executive KPI reporting with 99%+ accuracy.",
        ],
    )

    add_job(doc,
        "AT&T / YP Advertising", "Data Analyst", "Jan 2013 – Jan 2016",
        [
            "Built marketing campaign analytics and data governance solutions during AT&T's digital transformation "
            "from print-first to omni-channel advertising.",
            "Designed multi-channel attribution model reconciling digital (search, display, social) and print "
            "channels across 8,200+ advertiser accounts.",
            "Established enterprise data governance framework (data dictionary, stewardship processes) across "
            "6 source systems, reducing data quality incidents by 34% and improving billing accuracy.",
            "Performed advertiser segmentation and lifetime value modeling, driving 12% improvement in "
            "targeted renewal campaign conversion rates.",
            "Reduced manual reporting from 3-day turnaround to same-day delivery through automated dashboards.",
        ],
    )

    # Education
    add_section_heading(doc, "Education & Certifications")
    add_education(doc, EDU)

    path = os.path.join(OUTPUT_DIR, "Resume_PamelaAustin_DataAnalyst_BI.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


# ════════════════════════════════════════════════════════════════════════════
# RESUME 2 — Data Program Manager / Technical Project Manager
# ════════════════════════════════════════════════════════════════════════════

def build_resume_project_manager():
    doc = make_document()

    add_name_header(
        doc,
        "Pamela Austin",
        "Data Program Manager  |  Technical Project Manager",
        CONTACT,
    )

    add_section_heading(doc, "Professional Summary", space_before=0)
    add_summary(doc,
        "Strategic Data Program Manager with 13+ years of experience delivering complex analytics and "
        "technology initiatives across enterprise organizations (AT&T, Signet Jewelers, Syilum LLC). Skilled "
        "at scoping, prioritizing, and driving multi-workstream data programs from requirements to delivery. "
        "Track record of compressing timelines dramatically (6-week processes reduced to 5 days), managing "
        "200+ concurrent projects, eliminating bottlenecks for 150+ business users, and translating technical "
        "deliverables into C-suite narratives. Uniquely combines deep technical fluency (SQL, Python, BI "
        "platforms) with stakeholder management and cross-functional program coordination."
    )

    add_section_heading(doc, "Core Competencies")
    add_skills_table(doc, {
        "Program Management":   "Scope Definition, Milestone Planning, Risk Management, Delivery Oversight",
        "Stakeholder Mgmt":     "C-Suite Reporting, Executive Dashboards, 50+ Stakeholder Coordination",
        "Process Optimization": "Workflow Automation, Bottleneck Elimination, SLA Governance",
        "Technical Fluency":    "Power BI, SQL, Python, Snowflake, Databricks, ETL/ELT Pipelines",
        "Tools & Platforms":    "Workfront, Jira, Confluence, Salesforce, Fivetran, dbt, SAP",
        "Analytics & Reporting":"KPI Frameworks, Budget Variance, EBITDA Analysis, Campaign Reporting",
        "Data Governance":      "Policy Frameworks, Data Dictionaries, Quality Stewardship",
        "Business Domains":     "M&A Integration, Marketing Operations, Product Analytics, Finance",
    })

    add_section_heading(doc, "Professional Experience")

    add_job(doc,
        "Syilum LLC", "Freelance BI Program Lead", "Oct 2025 – Present",
        [
            "Manage end-to-end analytics program for a technology company: scope, prioritize, and deliver "
            "BI solutions aligned to product and executive roadmap priorities.",
            "Audited and consolidated 23 conflicting dashboards into a governed reporting framework with "
            "an agreed KPI dictionary—eliminating recurring stakeholder confusion and reporting delays.",
            "Managed analytics request backlog in Jira using sprint-based delivery cycles—maintaining "
            "a structured intake queue, acceptance criteria, and definition-of-done that reduced ad-hoc "
            "interruptions by 60% and improved delivery predictability.",
            "Coordinate directly with product management and executive stakeholders to define requirements; "
            "author Confluence specification pages and acceptance criteria to ensure full alignment before "
            "development begins, eliminating costly rework and scope creep.",
        ],
    )

    add_job(doc,
        "Signet Jewelers", "BI Analyst – M&A Program Support", "Jun 2025 – Sep 2025",
        [
            "Sole BI resource embedded in Corporate Development team managing $450M+ acquisition pipeline; "
            "coordinated reporting deliverables across Finance, Legal, and Operations under deal-timeline pressure.",
            "Managed financial data program compressing the financial consolidation process from 6 weeks "
            "to 5 days (92% improvement) through automated SQL reporting pipelines.",
            "Coordinated row-level security (RLS) governance framework for 50+ stakeholders with tiered "
            "access aligned to deal confidentiality requirements.",
            "Identified and escalated $12.6M in accounting discrepancies—including a $3.2M price renegotiation "
            "on a single deal—through automated SQL reconciliation, directly influencing acquisition decisions.",
            "Produced executive-ready Power BI reports and PowerPoint narratives for C-suite and Board-level "
            "deal review sessions.",
            "Managed scope and delivery of star schema design (12 fact tables, 18 dimensions) across a "
            "compressed engagement under strict deal timelines.",
        ],
    )

    add_job(doc,
        "AT&T", "Senior BI Analyst / Analytics Program Lead", "Mar 2024 – Mar 2025",
        [
            "Led delivery of AT&T Marketing Operations analytics platform—scoped and built 25+ Power BI "
            "dashboards tracking campaign performance, workforce costs, and budget for leadership.",
            "Managed concurrent analytics work across 200+ active campaigns using Workfront for campaign-level "
            "milestone and resource tracking alongside Jira for sprint planning, issue triage, and analytics "
            "delivery cycles; primary coordination point between project managers, Salesforce admins, and "
            "marketing operations directors.",
            "Drove self-service adoption program onboarding 150+ business users; authored Confluence data "
            "dictionaries, dashboard user guides, and onboarding runbooks that reduced IT ticket volume by 60% "
            "and eliminated manual data retrieval for 40% of recurring requests.",
            "Coordinated data integration program across Workfront, Salesforce, ZoomInfo, and Snowflake via "
            "Fivetran—defining data contracts, SLAs, and refresh cadence with technical and business teams.",
            "Managed predictive analytics initiative: scoped requirements, coordinated data access, trained "
            "and deployed ML forecasting model (89% accuracy), and delivered results to 150+ end users.",
            "Reduced project overruns by 28% and leadership escalations by 30% by embedding ML-generated "
            "at-risk project flags into operational dashboards used by program managers daily.",
        ],
    )

    add_job(doc,
        "Syilum LLC", "Product Data Analyst / Analytics Lead", "Jan 2016 – Mar 2024",
        [
            "Managed 8-year analytics function for a technology company: requirements gathering, delivery "
            "planning, and stakeholder reporting across product, marketing, and executive teams.",
            "Built and maintained 100+ dashboards supporting 10M+ daily events; managed analytics infrastructure "
            "serving product managers, engineering, and executive leadership.",
            "Led A/B testing program—defined experiment scope, coordinated with engineering for instrumentation, "
            "analyzed results, and presented findings that drove 23% feature adoption improvements.",
            "Established Jira-based analytics intake and sprint delivery process paired with a Confluence "
            "documentation library (data dictionaries, runbooks, stakeholder guides), reducing ad-hoc request "
            "volume by 60% through structured prioritization and self-service analytics.",
        ],
    )

    add_job(doc,
        "AT&T / YP Advertising", "Data Analyst", "Jan 2013 – Jan 2016",
        [
            "Managed data governance program across 6 source systems for 8,200+ advertiser accounts during "
            "AT&T's digital transformation—defined policies, assigned stewardship, and tracked resolution.",
            "Led advertiser data reconciliation initiative identifying 340 mis-classified active accounts, "
            "preventing a damaging win-back campaign and triggering an enterprise-wide data governance program.",
            "Coordinated multi-channel attribution project aligning digital and print campaign tracking, "
            "establishing unified ROI reporting across previously siloed tools.",
            "Delivered same-day automated campaign reporting replacing 3-day manual turnaround; managed "
            "stakeholder expectations and rollout across the sales and marketing division.",
        ],
    )

    add_section_heading(doc, "Education & Certifications")
    add_education(doc, EDU)

    path = os.path.join(OUTPUT_DIR, "Resume_PamelaAustin_DataProgramManager.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


# ════════════════════════════════════════════════════════════════════════════
# RESUME 3 — Analytics Engineer / Data Engineer
# ════════════════════════════════════════════════════════════════════════════

def build_resume_analytics_engineer():
    doc = make_document()

    add_name_header(
        doc,
        "Pamela Austin",
        "Analytics Engineer  |  Data Engineer",
        CONTACT,
    )

    add_section_heading(doc, "Professional Summary", space_before=0)
    add_summary(doc,
        "Analytics Engineer with 13+ years of hands-on experience designing and delivering data infrastructure "
        "at enterprise scale—ETL/ELT pipelines, cloud data warehouses, dimensional models, and ML systems. "
        "Deep expertise in Snowflake, Databricks, dbt, PySpark, and Python. Engineered cloud migration "
        "pipelines delivering 87.5% faster processing and $250K annual savings; deduplicated 150M+ records "
        "with zero data loss; designed star schema models supporting $450M+ M&A analytics; and built ML "
        "forecasting systems at 89–96.5% accuracy. Bridges the gap between raw data infrastructure and "
        "the analytical products that business stakeholders rely on."
    )

    add_section_heading(doc, "Technical Skills")
    add_skills_table(doc, {
        "Cloud Platforms":      "Snowflake, AWS, Databricks, Azure (familiarity)",
        "Languages":            "Python (Pandas, NumPy, scikit-learn, PySpark), SQL (Advanced), DAX",
        "Data Pipelines":       "ETL/ELT Design, Fivetran, dbt, Apache Spark, Automated Workflows",
        "Data Modeling":        "Star Schema, Dimensional Modeling, Fact/Dimension Design (12F/18D)",
        "ML & Analytics":       "Predictive Modeling, Forecasting, Anomaly Detection, A/B Testing",
        "BI & Visualization":   "Power BI, Tableau (output layer on engineered data)",
        "Data Quality":         "Validation Frameworks, Reconciliation Logic, Data Governance",
        "Integrations":         "Salesforce, Workfront, ZoomInfo, SAP via Fivetran/API",
    })

    add_section_heading(doc, "Professional Experience")

    add_job(doc,
        "Syilum LLC", "Freelance Analytics Engineer / BI Analyst", "Oct 2025 – Present",
        [
            "Design and maintain SQL/Snowflake analytics layer supporting product KPI dashboards and "
            "executive reporting; enforce 99%+ data accuracy via validation frameworks.",
            "Resolve semantic layer conflicts across 23 divergent dashboards by establishing a governed "
            "KPI dictionary—a canonical definitions layer consumed by all downstream reports.",
            "Build conversion funnel and user behavior SQL queries on Snowflake analyzing 10M+ daily "
            "events for product roadmap prioritization.",
            "Architect self-service analytics layer reducing ad-hoc engineering interruptions by 60%.",
        ],
    )

    add_job(doc,
        "Signet Jewelers", "BI & Data Analyst – M&A Analytics Engineering", "Jun 2025 – Sep 2025",
        [
            "Designed enterprise star schema (12 fact tables, 18 dimensions) for M&A financial analytics "
            "on Databricks/Snowflake, supporting 15.2M+ financial record processing.",
            "Engineered automated SQL reconciliation pipelines detecting $12.6M in accounting discrepancies "
            "by cross-validating acquisition target financials against SAP source data.",
            "Built Databricks-based ETL workflows automating financial consolidation, reducing processing "
            "time from 6 weeks to 5 days (92% improvement).",
            "Optimized Power BI DAX measures and Snowflake query plans, improving dashboard load times by "
            "65% and enabling drill-through analysis across acquisition scenarios.",
            "Implemented row-level security (RLS) framework in Power BI aligned to deal-level confidentiality "
            "tiers for 50+ stakeholders.",
        ],
    )

    add_job(doc,
        "AT&T", "Senior Analytics Engineer / BI Analyst", "Mar 2024 – Mar 2025",
        [
            "Engineered unified Snowflake reporting layer ingesting Workfront, Salesforce, and ZoomInfo "
            "via Fivetran; built dbt transformation models as conformed dimensions for the analytics layer.",
            "Optimized Snowflake SQL and Power BI data models, reducing dashboard query times from 45 "
            "seconds to 3 seconds (93% improvement) across 25+ executive-facing dashboards.",
            "Developed end-to-end ML forecasting pipeline in Python (scikit-learn, Pandas, NumPy) on top "
            "of the dbt/Snowflake dimensional layer; achieved 89% timeline forecast accuracy.",
            "Architected feature engineering pipeline consuming conformed Snowflake dimensions (employee, "
            "campaign, time hierarchies) for ML model training on 200+ campaign project histories.",
            "Deployed ML model outputs back to Snowflake and surfaced at-risk project flags in Power BI "
            "operational dashboards, reducing overruns 28% and escalations 30%.",
            "Built Fivetran data ingestion pipelines with dbt transformation layer for 150+ self-service users.",
        ],
    )

    add_job(doc,
        "Syilum LLC", "Product Data Analyst / Data Engineer", "Jan 2016 – Mar 2024",
        [
            "Built and maintained analytics data infrastructure supporting 10M+ daily events and 100+ dashboards "
            "over an 8-year engagement.",
            "Engineered ETL pipelines and dimensional models enabling consistent product KPI reporting across "
            "product, marketing, and executive stakeholders.",
            "Designed experiment infrastructure for A/B testing program; built event instrumentation specs "
            "and statistical analysis pipelines driving 23% feature adoption improvements.",
            "Deduplicated 150M+ records with zero data loss through Python/SQL deduplication pipelines, "
            "delivering 87.5% faster cloud processing and $250K annual savings.",
        ],
    )

    add_job(doc,
        "AT&T / YP Advertising", "Data Analyst", "Jan 2013 – Jan 2016",
        [
            "Built multi-channel attribution data model reconciling digital (search, display, social) and "
            "print advertising data across 6 source systems for 8,200+ advertiser accounts.",
            "Designed and implemented enterprise data governance framework: data dictionary, master advertiser "
            "record, and stewardship processes reducing data quality incidents by 34%.",
            "Engineered advertiser segmentation and lifetime value models driving 12% renewal campaign uplift.",
            "Automated campaign reporting pipeline reducing delivery from 3 days to same-day.",
        ],
    )

    # Personal Project
    add_section_heading(doc, "Notable Personal Project")
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(1)
    r1 = p.add_run("Exoplanet Habitability ML Classification  ")
    r1.bold = True
    r1.font.size = Pt(10)
    r1.font.color.rgb = BODY_BLACK
    r2 = p.add_run("(Python, scikit-learn, Pandas, Jupyter)")
    r2.font.size = Pt(9.5)
    r2.font.color.rgb = GRAY

    bp = doc.add_paragraph(style="List Bullet")
    bp.paragraph_format.left_indent  = Inches(0.2)
    bp.paragraph_format.space_before = Pt(1)
    bp.paragraph_format.space_after  = Pt(1)
    br = bp.add_run(
        "Built ML classification pipeline on NASA Kepler dataset (5,000+ exoplanets); achieved 96.5% accuracy "
        "predicting habitable world candidates using feature engineering, ensemble models, and cross-validation. "
        "Demonstrates applied ML skills beyond commercial project work."
    )
    br.font.size = Pt(9.5)
    br.font.name = "Calibri"

    add_section_heading(doc, "Education & Certifications")
    add_education(doc, EDU)

    path = os.path.join(OUTPUT_DIR, "Resume_PamelaAustin_AnalyticsEngineer.docx")
    doc.save(path)
    print(f"Saved: {path}")
    return path


# ── Main ─────────────────────────────────────────────────────────────────────

if __name__ == "__main__":
    print("Generating resumes...")
    build_resume_data_analyst()
    build_resume_project_manager()
    build_resume_analytics_engineer()
    print("\nAll three resumes generated successfully.")
